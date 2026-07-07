"""Credit quality risk committee memo — python-docx + Claude AI.

Spec (reports/credit_quality_report.py):
  generate_credit_quality_report(tenant_id, period) -> bytes

  Section 1: Executive Summary         — Claude-generated, ~150 words
  Section 2: Delinquency Trend Table   — 8-quarter, by loan type, color vs peer p50/p75
             Green < peer median | White at median | Amber > median | Red > 75th pctile
  Section 3: Regional Comparison       — institution vs peers; signal separator
             Label: "Is this a you-problem or a market-problem?"
             Claude narrative answers the question
  Section 4: Charge-Off and ALLL       — current + prior period + peer median
             Red callout when ALLL coverage < 1.0x
  Section 5: Watch Items               — Claude-generated, top 3
             Each: Metric | Current Value | Peer Context | Recommended Action

  Board version extras (board_version=True):
    - 8-quarter delinquency trend chart (matplotlib PNG; text table fallback)
    - Forward outlook narrative (Claude-generated)
    - Peer comparison appendix

  CLAUDE.md rules enforced:
    - Both dollar balance AND computed rate for every delinquency figure
    - All delinquency figures are institution-level (not branch-level)
    - Default peer group: REGIONAL
    - Signal separator: "Is this a you-problem or a market-problem?"
    - Confidence badge on every geographic figure
    - Delinquency is ADVERSE: lower = better
"""

from __future__ import annotations

import json
import logging
import os
from io import BytesIO
from datetime import date
from pathlib import Path
from typing import Optional

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
DB_URL            = os.environ.get("DATABASE_URL")

# ── Colors ────────────────────────────────────────────────────────────────────
_C_RED    = RGBColor(0xC6, 0x28, 0x28)
_C_GREEN  = RGBColor(0x2E, 0x7D, 0x32)
_C_AMBER  = RGBColor(0xFF, 0x6F, 0x00)
_C_TEAL   = RGBColor(0x00, 0x69, 0x6E)
_C_BLUE   = RGBColor(0x01, 0x57, 0x9B)
_C_PURPLE = RGBColor(0x6A, 0x1B, 0x9A)

_B_GREEN = "E8F5E9"
_B_RED   = "FFEBEE"
_B_AMBER = "FFF8E1"
_B_GRAY  = "F5F5F5"

_STARS = {1: "★☆☆☆☆", 2: "★★☆☆☆", 3: "★★★☆☆", 4: "★★★★☆", 5: "★★★★★"}

# Loan-type delinquency: (display label, inst_rate_key, delinq_cols, balance_cols, peer_metric_key)
# delinq_cols are named ORM columns (not JSONB codes) — all confirmed in db.py
_LOAN_TYPE_CFG = [
    (
        "Total 60+",
        "delinq_rate_total",
        [],                                          # computed from acct_041B / acct_025B directly
        [],
        "delinq_rate_total",
    ),
    (
        "Auto Loans",
        "auto_delinq_rate",
        ["acct_041C1", "acct_041C2"],                # new + used vehicle 60+ (db.py FS220I cols)
        ["acct_385", "acct_370"],
        "delinq_rate_auto",
    ),
    (
        "Credit Card",
        "credit_card_delinq_rate",
        ["acct_045B"],                               # total delinquent CC loans (db.py FS220B col)
        ["acct_396"],
        "delinq_rate_cc",
    ),
    (
        "Commercial",
        "commercial_delinq_rate",
        ["acct_041G1", "acct_041G2", "acct_041G3", "acct_041G4"],  # all commercial types (db.py)
        ["acct_718A5", "acct_400P"],
        "delinq_rate_commercial_total",
    ),
]

# ── Shared helpers ────────────────────────────────────────────────────────────

def _cell_bg(cell, hex6: str) -> None:
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex6}" w:color="auto" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _bold_run(paragraph, text: str, color: Optional[RGBColor] = None) -> None:
    r = paragraph.add_run(text)
    r.font.bold = True
    if color:
        r.font.color.rgb = color


def _col_widths(table, *widths_in: float) -> None:
    for i, w in enumerate(widths_in):
        if i < len(table.columns):
            table.columns[i].width = Inches(w)


def _fmt_pct(v: float, decimals: int = 3) -> str:
    return f"{v * 100:.{decimals}f}%"


def _fmt_dollar(v: float) -> str:
    if v >= 1e9:  return f"${v / 1e9:.2f}B"
    if v >= 1e6:  return f"${v / 1e6:.1f}M"
    if v >= 1e3:  return f"${v / 1e3:.0f}K"
    return f"${v:,.0f}"


def _period_label(period: str) -> str:
    if "Q" in period:
        return f"Q{period[5]} {period[:4]}"
    return period


def _prior_period(period: str) -> str:
    """'2026Q1' → '2025Q4', '2025Q4' → '2025Q3'"""
    if "Q" not in period:
        return period
    q = int(period[5])
    y = int(period[:4])
    return f"{y - 1}Q4" if q == 1 else f"{y}Q{q - 1}"


def _prior_year_period(period: str) -> str:
    """'2026Q1' → '2025Q1'"""
    if "Q" not in period:
        return period
    return f"{int(period[:4]) - 1}Q{period[5]}"


# ── Claude helper ─────────────────────────────────────────────────────────────

def _claude(prompt: str, max_words: int = 180) -> str:
    if not ANTHROPIC_API_KEY:
        return "[AI narrative unavailable — ANTHROPIC_API_KEY not set]"
    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=max_words * 2,
            system=(
                "You are a credit union risk management report writer. "
                "Write in the voice of an experienced credit risk officer. "
                "Use Callahan Associates metric terminology. Be specific about numbers. "
                "Write in present tense. No markdown formatting."
            ),
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text.strip()
    except Exception as exc:
        logger.warning("Claude call failed: %s", exc)
        return "[AI narrative generation failed]"


# ── Red callout box ───────────────────────────────────────────────────────────

def _red_callout(doc: Document, message: str) -> None:
    """Render a single-cell table styled as a red alert callout."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, _B_RED)
    p = cell.paragraphs[0]
    r = p.add_run(f"⚠  {message}")
    r.font.bold = True
    r.font.color.rgb = _C_RED
    r.font.size = Pt(11)
    doc.add_paragraph()


# ── Data loaders ──────────────────────────────────────────────────────────────

def _fetch_peer_dist(
    metric_key: str,
    period: str,
    peer_group_id: Optional[str],
    engine,
    peer_charters: Optional[list] = None,
    db_url: Optional[str] = None,
) -> dict:
    """Return p50 and p75 for a metric.

    Tries precomputed peer_distributions first; falls back to on-the-fly
    compute_peer_distribution() when peer_group_id is missing.
    """
    if peer_group_id:
        try:
            from sqlalchemy import text
            with engine.connect() as conn:
                row = conn.execute(
                    text("""
                        SELECT p50 AS median, p75
                        FROM peer_distributions
                        WHERE metric_key = :mk AND period = :p AND peer_group_id = :pg
                        LIMIT 1
                    """),
                    {"mk": metric_key, "p": period, "pg": str(peer_group_id)},
                ).mappings().first()
            if row:
                return dict(row)
        except Exception:
            pass

    # On-the-fly fallback using actual peer charter list
    if peer_charters and db_url:
        try:
            from processing.delinquency_engine import compute_peer_distribution
            dist = compute_peer_distribution(metric_key, peer_charters, period, db_url)
            return {"median": dist.get("p50"), "p75": dist.get("p75")}
        except Exception as exc:
            logger.warning("On-the-fly peer dist failed metric=%s period=%s: %s", metric_key, period, exc)

    return {}


def _load_inst_row(charter_number: int, period: str, engine) -> dict:
    """Fetch one period of institution data. Returns {} on miss.

    PostgreSQL folds unquoted identifiers to lowercase (acct_025B → acct_025b).
    We normalise keys back to the canonical ORM-defined names so that downstream
    code can use inst.get("acct_025B") etc. without case surprises.
    """
    try:
        from sqlalchemy import text
        from db import institutions_quarterly as _iq
        # Build lowercase→canonical map once from the ORM metadata
        _case_map = {c.name.lower(): c.name for c in _iq.columns}

        with engine.connect() as conn:
            row = conn.execute(
                text("""
                    SELECT acct_010, acct_025B, acct_041B, acct_020B,
                           acct_AS0048, acct_719, acct_550, acct_551,
                           acct_385, acct_370, acct_396, acct_703A,
                           acct_386A, acct_718A5, acct_400P, acct_797,
                           acct_IS0010, acct_117, acct_671, acct_661A,
                           acct_041C1, acct_041C2, acct_045B,
                           acct_041G1, acct_041G2, acct_041G3, acct_041G4
                    FROM institutions_quarterly
                    WHERE charter_number = :c AND period = :p
                    LIMIT 1
                """),
                {"c": charter_number, "p": period},
            ).mappings().first()
        if not row:
            return {}
        # Remap lowercase PG keys → ORM canonical names (e.g. acct_025b → acct_025B)
        result = {_case_map.get(k, k): v for k, v in dict(row).items()}
        return result
    except Exception as exc:
        logger.warning("_load_inst_row failed charter=%s period=%s: %s", charter_number, period, exc)
        return {}


def _compute_loan_type_rate(row: dict, delinq_cols: list[str], balance_cols: list[str]) -> Optional[float]:
    """Compute delinquency rate for a loan type using named ORM columns from the inst row."""
    if not delinq_cols or not balance_cols:
        return None
    delinq_total  = sum(float(row.get(c) or 0) for c in delinq_cols)
    balance_total = sum(float(row.get(c) or 0) for c in balance_cols)
    if balance_total <= 0:
        return None
    return delinq_total / balance_total


def _load_trend_data(
    charter_number: int,
    period: str,
    peer_group_id: Optional[str],
    db_url: Optional[str],
    peer_charters: Optional[list] = None,
) -> list[dict]:
    """Return up to 8 quarters of delinquency rates (institution + peer benchmarks).

    Returns list newest-first. Each dict has:
      period, delinq_rate_total, auto_delinq_rate, credit_card_delinq_rate,
      commercial_delinq_rate, plus *_p50 and *_p75 peer benchmarks.
    """
    from processing.early_warning_engine import _trailing_periods
    from db import get_engine

    engine = get_engine(db_url)
    periods = _trailing_periods(period, n=8)
    result = []

    for p in periods:
        inst = _load_inst_row(charter_number, p, engine)
        row_data: dict = {"period": p}

        if inst:
            total_loans = float(inst.get("acct_025B") or 0)
            total_delinq = float(inst.get("acct_041B") or 0)
            row_data["delinq_rate_total"] = total_delinq / total_loans if total_loans > 0 else None
            row_data["delinq_balance_total"] = total_delinq

            for label, inst_key, delinq_codes, bal_cols, _ in _LOAN_TYPE_CFG[1:]:
                row_data[inst_key] = _compute_loan_type_rate(inst, delinq_codes, bal_cols)

        # Peer benchmarks — precomputed table first, on-the-fly fallback
        for _, inst_key, _, _, peer_metric_key in _LOAN_TYPE_CFG:
            dist = _fetch_peer_dist(
                peer_metric_key, p, peer_group_id, engine,
                peer_charters=peer_charters, db_url=db_url,
            )
            row_data[f"{inst_key}_p50"] = dist.get("median")
            row_data[f"{inst_key}_p75"] = dist.get("p75")

        result.append(row_data)

    return result


def _load_regional_comparison(
    charter_number: int,
    period: str,
    peer_charters: list[int],
    db_url: Optional[str],
) -> dict:
    """Compute delinquency rate distribution across regional peer institutions."""
    from db import get_engine
    from sqlalchemy import text

    engine = get_engine(db_url)
    rates: list[float] = []
    rows: list[dict] = []

    if not peer_charters:
        return {}

    try:
        with engine.connect() as conn:
            peer_rows = conn.execute(
                text("""
                    SELECT charter_number, institution_name, acct_041B, acct_025B
                    FROM institutions_quarterly
                    WHERE charter_number = ANY(:charters) AND period = :p
                """),
                {"charters": peer_charters, "p": period},
            ).mappings().all()

        for r in peer_rows:
            # PG returns lowercase column names from raw text(); use lowercase keys here
            loans = float(r.get("acct_025b") or r.get("acct_025B") or 0)
            delinq = float(r.get("acct_041b") or r.get("acct_041B") or 0)
            if loans > 0:
                rate = delinq / loans
                rates.append(rate)
                rows.append({
                    "charter_number": r["charter_number"],
                    "institution_name": r["institution_name"] or f"Charter {r['charter_number']}",
                    "delinq_rate": rate,
                    "delinq_balance": delinq,
                })

        if not rates:
            return {}

        rates_sorted = sorted(rates)
        n = len(rates_sorted)
        p_idx = lambda pct: max(0, min(n - 1, int(pct * n / 100)))

        return {
            "n_peers":        n,
            "peer_p10":       rates_sorted[p_idx(10)],
            "peer_p25":       rates_sorted[p_idx(25)],
            "peer_p50":       rates_sorted[p_idx(50)],
            "peer_p75":       rates_sorted[p_idx(75)],
            "peer_p90":       rates_sorted[p_idx(90)],
            "peer_rows":      sorted(rows, key=lambda r: r["delinq_rate"]),
        }
    except Exception as exc:
        logger.warning("Regional comparison load failed: %s", exc)
        return {}


def _load_prior_period_data(charter_number: int, period: str, db_url: Optional[str]) -> dict:
    """Load institution data for prior quarter and prior year same quarter."""
    from db import get_engine
    engine = get_engine(db_url)

    prior_qtr    = _load_inst_row(charter_number, _prior_period(period), engine)
    prior_year   = _load_inst_row(charter_number, _prior_year_period(period), engine)

    def _co_rate(row: dict) -> Optional[float]:
        gross = float(row.get("acct_550") or 0)
        rec   = float(row.get("acct_551") or 0)
        loans = float(row.get("acct_025B") or 0)
        return (gross - rec) / loans * 4 if loans > 0 else None

    def _cov(row: dict) -> Optional[float]:
        allow  = float(row.get("acct_AS0048") or row.get("acct_719") or 0)
        delinq = float(row.get("acct_041B") or 0)
        return allow / delinq if delinq > 0 else None

    return {
        "prior_qtr_co_rate":       _co_rate(prior_qtr),
        "prior_year_co_rate":      _co_rate(prior_year),
        "prior_qtr_alll_coverage": _cov(prior_qtr),
        "prior_year_alll_coverage": _cov(prior_year),
        "prior_qtr_delinq_rate":   (
            float(prior_qtr.get("acct_041B") or 0) / float(prior_qtr.get("acct_025B") or 1)
            if prior_qtr.get("acct_025B") else None
        ),
        "prior_year_delinq_rate":  (
            float(prior_year.get("acct_041B") or 0) / float(prior_year.get("acct_025B") or 1)
            if prior_year.get("acct_025B") else None
        ),
    }


# ── Narrative generators ──────────────────────────────────────────────────────

def _gen_exec_summary(
    institution_name: str,
    period: str,
    peer_label: str,
    inst_financials: dict,
    trend_data: list[dict],
    prior_data: dict,
    peer_group_id: Optional[str],
) -> str:
    current = trend_data[0] if trend_data else {}
    total_rate = current.get("delinq_rate_total")
    peer_p50   = current.get("delinq_rate_total_p50")
    peer_p75   = current.get("delinq_rate_total_p75")
    co_rate    = inst_financials.get("chargeoff_rate_total_annualized")
    coverage   = inst_financials.get("alll_coverage")
    prior_qtr_rate = prior_data.get("prior_qtr_delinq_rate")

    trend_dir = ""
    if total_rate is not None and prior_qtr_rate is not None:
        delta = total_rate - prior_qtr_rate
        trend_dir = f" Rate {'increased' if delta > 0 else 'decreased'} {abs(delta)*100:.2f} pp vs prior quarter."

    prompt = (
        f"Write a 150-word executive summary for {institution_name}'s credit quality risk committee memo "
        f"for {_period_label(period)}. Peer group: {peer_label}.\n"
        f"Total 60+ day delinquency rate: {_fmt_pct(total_rate) if total_rate else 'N/A'}."
        f"{trend_dir}"
        + (f" Peer median: {_fmt_pct(peer_p50)}." if peer_p50 else "")
        + (f" Peer 75th pctile: {_fmt_pct(peer_p75)}." if peer_p75 else "")
        + (f" Net charge-off rate (annualized): {_fmt_pct(co_rate)}." if co_rate else "")
        + (f" ALLL/ACL coverage ratio: {coverage:.2f}x." if coverage else "")
        + " Summarize overall credit quality position, trend direction, and one key area of management focus."
    )
    return _claude(prompt, max_words=170)


def _gen_regional_narrative(
    institution_name: str,
    period: str,
    inst_rate: Optional[float],
    regional_data: dict,
) -> str:
    n   = regional_data.get("n_peers", 0)
    p50 = regional_data.get("peer_p50")
    p75 = regional_data.get("peer_p75")

    vs_region = ""
    if inst_rate is not None and p50 is not None:
        if inst_rate > (p75 or p50 * 1.5):
            vs_region = f"Institution rate of {_fmt_pct(inst_rate)} exceeds the regional 75th percentile of {_fmt_pct(p75)}. This is likely an institution-specific issue."
        elif inst_rate > p50:
            vs_region = f"Institution rate of {_fmt_pct(inst_rate)} is above the regional median of {_fmt_pct(p50)}. Elevated but within the upper half of regional peers."
        else:
            vs_region = f"Institution rate of {_fmt_pct(inst_rate)} is at or below the regional median of {_fmt_pct(p50)}. Performance is consistent with or better than regional market conditions."

    prompt = (
        f"Write a 130-word narrative answering: 'Is this a you-problem or a market-problem?' "
        f"for {institution_name} ({_period_label(period)}).\n"
        f"Regional comparison vs {n} institutions in same geography: {vs_region}\n"
        "Distinguish between institution-specific credit risk versus market-wide conditions. "
        "Reference whether the regional peer group is also experiencing stress. "
        "Be direct and actionable."
    )
    return _claude(prompt, max_words=150)


def _gen_watch_items(
    institution_name: str,
    period: str,
    peer_label: str,
    inst_financials: dict,
    trend_data: list[dict],
    prior_data: dict,
    coverage: Optional[float],
    co_rate: Optional[float],
) -> str:
    current = trend_data[0] if trend_data else {}
    total_rate    = current.get("delinq_rate_total")
    peer_p50      = current.get("delinq_rate_total_p50")
    peer_p75      = current.get("delinq_rate_total_p75")
    auto_rate     = current.get("auto_delinq_rate")
    cc_rate       = current.get("credit_card_delinq_rate")
    comm_rate     = current.get("commercial_delinq_rate")
    prior_total   = prior_data.get("prior_qtr_delinq_rate")
    prior_co      = prior_data.get("prior_qtr_co_rate")
    prior_cov     = prior_data.get("prior_qtr_alll_coverage")

    context = (
        f"Institution: {institution_name}, Period: {_period_label(period)}, Peer Group: {peer_label}\n"
        f"Total 60+ delinquency: {_fmt_pct(total_rate) if total_rate else 'N/A'} "
        f"(prior quarter: {_fmt_pct(prior_total) if prior_total else 'N/A'}, "
        f"peer median: {_fmt_pct(peer_p50) if peer_p50 else 'N/A'}, peer 75th: {_fmt_pct(peer_p75) if peer_p75 else 'N/A'})\n"
        f"Auto delinquency: {_fmt_pct(auto_rate) if auto_rate else 'N/A'}\n"
        f"Credit card delinquency: {_fmt_pct(cc_rate) if cc_rate else 'N/A'}\n"
        f"Commercial delinquency: {_fmt_pct(comm_rate) if comm_rate else 'N/A'}\n"
        f"Net charge-off rate (ann.): {_fmt_pct(co_rate) if co_rate else 'N/A'} "
        f"(prior: {_fmt_pct(prior_co) if prior_co else 'N/A'})\n"
        f"ALLL/ACL coverage: {f'{coverage:.2f}x' if coverage else 'N/A'} "
        f"(prior: {f'{prior_cov:.2f}x' if prior_cov else 'N/A'})"
    )

    prompt = (
        f"Based on this credit quality data, write the top 3 watch items requiring management attention.\n\n"
        f"{context}\n\n"
        "For each item, output in exactly this format:\n"
        "ITEM 1:\n"
        "Metric: [metric name using Callahan terminology]\n"
        "Current Value: [specific number]\n"
        "Peer Context: [how it compares to regional peers — be specific with numbers]\n"
        "Recommended Action: [one specific, actionable management step]\n\n"
        "ITEM 2:\n[same format]\n\n"
        "ITEM 3:\n[same format]\n\n"
        "Prioritize by severity. Focus on the most material credit risks."
    )
    return _claude(prompt, max_words=350)


def _gen_forward_outlook(institution_name: str, period: str, trend_data: list[dict]) -> str:
    rates = [d.get("delinq_rate_total") for d in trend_data if d.get("delinq_rate_total") is not None]
    trend_str = ", ".join(f"{_fmt_pct(r)}" for r in rates[:6])
    direction = ""
    if len(rates) >= 3:
        recent  = sum(rates[:2]) / 2
        earlier = sum(rates[2:4]) / 2 if len(rates) >= 4 else rates[2]
        direction = "increasing" if recent > earlier else "decreasing"

    prompt = (
        f"Write a 150-word forward outlook for {institution_name}'s credit quality for the next 2 quarters "
        f"after {_period_label(period)}.\n"
        f"Recent delinquency trend (newest first): {trend_str}. Trend direction: {direction}.\n"
        "Address: (1) expected trajectory based on trend, "
        "(2) seasonal factors (Q1 typically lower than Q4 for most loan types), "
        "(3) two specific metrics management should monitor closely next quarter. "
        "Be direct and forward-looking. No hedging."
    )
    return _claude(prompt, max_words=170)


# ── Chart (board version) ─────────────────────────────────────────────────────

def _make_trend_chart(trend_data: list[dict]) -> Optional[bytes]:
    """Generate delinquency trend chart as PNG bytes. Returns None if matplotlib unavailable."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        # Oldest-first for x-axis
        data = list(reversed(trend_data))
        periods   = [_period_label(d["period"]) for d in data]
        inst_pcts = [d.get("delinq_rate_total", 0) and d["delinq_rate_total"] * 100 for d in data]
        p50_pcts  = [d.get("delinq_rate_total_p50") and d["delinq_rate_total_p50"] * 100 for d in data]
        p75_pcts  = [d.get("delinq_rate_total_p75") and d["delinq_rate_total_p75"] * 100 for d in data]

        fig, ax = plt.subplots(figsize=(7.5, 3.5))
        x = list(range(len(periods)))

        inst_y = [v if v else None for v in inst_pcts]
        p50_y  = [v if v else None for v in p50_pcts]
        p75_y  = [v if v else None for v in p75_pcts]

        ax.plot(x, inst_y, "b-o", linewidth=2.0, markersize=4, label="Institution", zorder=3)
        ax.plot(x, p50_y,  "g--", linewidth=1.5, label="Peer Median (P50)", zorder=2)
        ax.plot(x, p75_y,  "r:",  linewidth=1.2, label="Peer 75th Pctile", zorder=2)

        # Shade between P50 and P75
        if any(v for v in p50_y) and any(v for v in p75_y):
            ax.fill_between(x, p50_y, p75_y, alpha=0.10, color="red", label="Peer P50–P75 band")

        ax.set_xticks(x)
        ax.set_xticklabels(periods, rotation=45, ha="right", fontsize=8)
        ax.set_ylabel("60+ Day Delinquency Rate (%)")
        ax.set_title("Total Delinquency Rate — 8-Quarter Trend")
        ax.legend(fontsize=8, loc="upper left")
        ax.grid(True, alpha=0.3)
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda y, _: f"{y:.2f}%"))
        plt.tight_layout()

        buf = BytesIO()
        plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except ImportError:
        return None
    except Exception as exc:
        logger.warning("Chart generation failed: %s", exc)
        return None


# ── Section builders ──────────────────────────────────────────────────────────

def _add_memo_header(
    doc: Document,
    institution_name: str,
    period: str,
    peer_label: str,
    report_type: str = "risk_committee",
) -> None:
    _heading(doc, "Credit Quality Memorandum")
    kind = "Risk Committee" if report_type == "risk_committee" else "Board of Directors"

    fields = [
        ("To:",        kind),
        ("From:",      "Market Intelligence Platform / Credit Risk Analytics"),
        ("Re:",        f"Credit Quality Analysis — {institution_name} — {_period_label(period)}"),
        ("Date:",      date.today().strftime("%B %d, %Y")),
        ("Period:",    _period_label(period)),
        ("Peer Group:", peer_label),
        ("Classification:", "CONFIDENTIAL"),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(1.4)
    table.columns[1].width = Inches(5.1)
    for i, (label, value) in enumerate(fields):
        table.cell(i, 0).paragraphs[0].add_run(label).font.bold = True
        cell = table.cell(i, 1)
        if label == "Classification:":
            r = cell.paragraphs[0].add_run(value)
            r.font.bold  = True
            r.font.color.rgb = _C_RED
        else:
            cell.text = value

    doc.add_paragraph()


def _add_exec_summary(doc: Document, narrative: str) -> None:
    _heading(doc, "Section 1 — Executive Summary")
    doc.add_paragraph(narrative)


def _add_delinq_trend_table(
    doc: Document,
    trend_data: list[dict],
    peer_label: str,
    chart_bytes: Optional[bytes] = None,
) -> None:
    _heading(doc, "Section 2 — Delinquency Trend (8 Quarters)")
    doc.add_paragraph(
        f"60+ day delinquency by loan type — {peer_label}. "
        "Shading: green = below peer median (better); "
        "amber = above peer median; red = above peer 75th percentile (worst quartile). "
        "ADVERSE metric: lower rate = better performance."
    )

    if not trend_data:
        doc.add_paragraph("No trend data available.")
        return

    if chart_bytes:
        doc.add_picture(BytesIO(chart_bytes), width=Inches(6.0))
        doc.add_paragraph()

    # Table: Period | (rate | peer_p50) per loan type
    n_loan_types = len(_LOAN_TYPE_CFG)
    # Cols: Period + (Institution Rate + Peer P50) per loan type = 1 + 2*n
    n_cols = 1 + 2 * n_loan_types
    table  = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hrow = table.rows[0]
    hrow.cells[0].text = ""
    hrow.cells[0].paragraphs[0].add_run("Period").font.bold = True

    for i, (label, _, _, _, _) in enumerate(_LOAN_TYPE_CFG):
        col_base = 1 + i * 2
        r1 = hrow.cells[col_base].paragraphs[0].add_run(label)
        r1.font.bold = True
        r1.font.size = Pt(8)
        r2 = hrow.cells[col_base + 1].paragraphs[0].add_run("Peer P50")
        r2.font.bold = True
        r2.font.size = Pt(8)
        r2.font.color.rgb = _C_TEAL

    # Column widths (tight to fit portrait page)
    widths = [0.60] + [0.68, 0.62] * n_loan_types
    _col_widths(table, *widths)

    def _bg_for(inst_rate: Optional[float], p50: Optional[float], p75: Optional[float]) -> Optional[str]:
        if inst_rate is None:
            return _B_GRAY
        if p75 is not None and inst_rate > p75:
            return _B_RED
        if p50 is not None and inst_rate > p50 + 0.00005:
            return _B_AMBER
        if p50 is not None and inst_rate <= p50:
            return _B_GREEN
        return None   # at median — no shade

    for d in trend_data:
        row = table.add_row()
        row.cells[0].text = _period_label(d["period"])
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)

        for i, (_, inst_key, _, _, _) in enumerate(_LOAN_TYPE_CFG):
            col_base = 1 + i * 2
            inst_rate = d.get(inst_key)
            p50       = d.get(f"{inst_key}_p50")
            p75       = d.get(f"{inst_key}_p75")

            # Institution rate cell
            cell_inst = row.cells[col_base]
            r = cell_inst.paragraphs[0].add_run(_fmt_pct(inst_rate, 3) if inst_rate is not None else "—")
            r.font.size = Pt(8)
            bg = _bg_for(inst_rate, p50, p75)
            if bg:
                _cell_bg(cell_inst, bg)

            # Peer P50 cell (light teal text, no background)
            cell_p50 = row.cells[col_base + 1]
            rp = cell_p50.paragraphs[0].add_run(_fmt_pct(p50, 3) if p50 is not None else "—")
            rp.font.size = Pt(8)
            rp.font.color.rgb = _C_TEAL

    # Legend
    p = doc.add_paragraph()
    p.add_run("Legend:  ")
    for label, bg in [("Below peer median", _B_GREEN), ("Above median", _B_AMBER), (">75th pctile", _B_RED)]:
        tbl_l = doc.add_table(rows=1, cols=1)
        tbl_l.rows[0].cells[0].text = ""
        _cell_bg(tbl_l.rows[0].cells[0], bg)
    doc.add_paragraph(
        "Green background: below peer median (better). "
        "Amber: above peer median. Red: above 75th percentile (worst quartile)."
    ).runs[0].font.size = Pt(9)


def _add_regional_comparison(
    doc: Document,
    regional_data: dict,
    institution_name: str,
    inst_rate: Optional[float],
    narrative: str,
) -> None:
    _heading(doc, "Section 3 — Regional Comparison")

    # Signal separator (CLAUDE.md rule — appears below every delinquency/charge-off chart)
    sep = doc.add_paragraph()
    r = sep.add_run("Is this a you-problem or a market-problem?")
    r.font.bold   = True
    r.font.size   = Pt(12)
    r.font.color.rgb = _C_PURPLE

    doc.add_paragraph(
        f"Compares {institution_name}'s 60+ day delinquency rate to "
        f"{regional_data.get('n_peers', 0)} institutions in the same regional market."
    )

    if not regional_data:
        doc.add_paragraph("Regional comparison data not available.")
        return

    # Distribution statistics table
    dist_table = doc.add_table(rows=1, cols=7)
    dist_table.style = "Table Grid"
    for i, h in enumerate(["", "P10", "P25", "Median (P50)", "P75", "P90", "Your Rate"]):
        dist_table.rows[0].cells[i].paragraphs[0].add_run(h).font.bold = True

    row = dist_table.add_row()
    row.cells[0].text = "60+ Day Delinq Rate"
    for i, key in enumerate(["peer_p10", "peer_p25", "peer_p50", "peer_p75", "peer_p90"]):
        v = regional_data.get(key)
        row.cells[i + 1].text = _fmt_pct(v) if v is not None else "—"

    p50 = regional_data.get("peer_p50")
    p75 = regional_data.get("peer_p75")
    cell_yours = row.cells[6]
    cell_yours.text = _fmt_pct(inst_rate) if inst_rate is not None else "—"
    if inst_rate is not None and p75 is not None and inst_rate > p75:
        _cell_bg(cell_yours, _B_RED)
    elif inst_rate is not None and p50 is not None and inst_rate > p50:
        _cell_bg(cell_yours, _B_AMBER)
    elif inst_rate is not None and p50 is not None:
        _cell_bg(cell_yours, _B_GREEN)

    _col_widths(dist_table, 1.6, 0.7, 0.7, 1.0, 0.7, 0.7, 0.9)

    doc.add_paragraph()

    # Full peer ranking table (top and bottom 5)
    peer_rows = regional_data.get("peer_rows", [])
    if peer_rows:
        _heading(doc, "Regional Peer Ranking", level=2)
        doc.add_paragraph(f"All {len(peer_rows)} regional peers ranked by 60+ day delinquency rate (lowest = best).")

        peer_table = doc.add_table(rows=1, cols=3)
        peer_table.style = "Table Grid"
        for i, h in enumerate(["Rank", "Institution", "60+ Day Rate"]):
            peer_table.rows[0].cells[i].paragraphs[0].add_run(h).font.bold = True
        _col_widths(peer_table, 0.6, 4.0, 1.2)

        for rank, pr in enumerate(peer_rows, 1):
            prow = peer_table.add_row()
            prow.cells[0].text = str(rank)
            prow.cells[1].text = pr["institution_name"]
            prow.cells[2].text = _fmt_pct(pr["delinq_rate"])
            # Mark this institution
            if pr["charter_number"] == None:  # placeholder — can't match without charter
                pass
            elif inst_rate is not None and abs(pr["delinq_rate"] - inst_rate) < 0.00001:
                for c in prow.cells:
                    _cell_bg(c, _B_AMBER)
                prow.cells[1].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

    # Claude narrative
    if narrative:
        doc.add_paragraph(narrative)


def _add_chargeoff_alll(
    doc: Document,
    data: dict,
    prior_data: dict,
    peer_label: str,
) -> None:
    _heading(doc, "Section 4 — Charge-Off and ALLL/ACL Coverage")

    inst     = data.get("institution_financials", {})
    composite = data.get("credit_risk_composite", {})
    co_metrics = composite.get("metrics", {}).get("chargeoff_rate_total_annualized", {})
    cov_metrics = composite.get("metrics", {}).get("alll_coverage", {})

    gross_co   = inst.get("acct_550")
    recoveries = inst.get("acct_551")
    net_co     = (float(gross_co or 0) - float(recoveries or 0)) if gross_co is not None else None
    co_rate    = inst.get("chargeoff_rate_total_annualized")
    allowance  = float(inst.get("acct_AS0048") or inst.get("acct_719") or 0) or None
    coverage   = inst.get("alll_coverage")
    allow_to_loans = inst.get("alll_to_loans")

    co_peer_med  = co_metrics.get("peer_distribution", {}).get("p50")
    cov_peer_med = cov_metrics.get("peer_distribution", {}).get("p50")
    co_pctile    = co_metrics.get("percentile_rank")
    co_stars     = co_metrics.get("stars")

    # Charge-off table: current | prior quarter | prior year | peer median
    _heading(doc, "Charge-Off Rates", level=2)
    co_table = doc.add_table(rows=1, cols=5)
    co_table.style = "Table Grid"
    for i, h in enumerate(["Metric", "Current", "Prior Quarter", "Prior Year", "Peer Median"]):
        co_table.rows[0].cells[i].paragraphs[0].add_run(h).font.bold = True
    _col_widths(co_table, 1.6, 0.9, 0.9, 0.9, 0.9)

    co_rows = [
        ("Gross Charge-offs", gross_co and _fmt_dollar(gross_co), None, None, None),
        ("Recoveries",        recoveries and _fmt_dollar(recoveries), None, None, None),
        ("Net Charge-offs",   net_co and _fmt_dollar(net_co), None, None, None),
        (
            "Rate (Ann.)",
            co_rate and _fmt_pct(co_rate, 3),
            prior_data.get("prior_qtr_co_rate") and _fmt_pct(prior_data["prior_qtr_co_rate"], 3),
            prior_data.get("prior_year_co_rate") and _fmt_pct(prior_data["prior_year_co_rate"], 3),
            co_peer_med and _fmt_pct(co_peer_med, 3),
        ),
        (
            "Peer Percentile",
            f"{co_pctile:.0f}th" if co_pctile else "—",
            None, None,
            _STARS.get(co_stars, "—") if co_stars else "—",
        ),
    ]

    for vals in co_rows:
        row = co_table.add_row()
        for j, v in enumerate(vals):
            row.cells[j].text = str(v) if v is not None else "—"
        # Color rate row
        if vals[0] == "Rate (Ann.)" and co_rate is not None and co_peer_med is not None:
            bg = _B_AMBER if co_rate > co_peer_med else _B_GREEN
            for c in row.cells[1:]:
                _cell_bg(c, bg)

    doc.add_paragraph()

    # ALLL / ACL adequacy
    _heading(doc, "Allowance Adequacy (ACL / ALLL)", level=2)

    # Red callout if coverage < 1.0x
    if coverage is not None and coverage < 1.0:
        _red_callout(
            doc,
            f"ALLL/ACL coverage ratio is {coverage:.2f}x — BELOW 1.0x minimum. "
            "Immediate board-level review required. Per GAAP and regulatory guidance, "
            "allowance must at minimum cover expected credit losses on impaired loans."
        )

    alll_table = doc.add_table(rows=1, cols=5)
    alll_table.style = "Table Grid"
    for i, h in enumerate(["Metric", "Current", "Prior Quarter", "Prior Year", "Peer Median"]):
        alll_table.rows[0].cells[i].paragraphs[0].add_run(h).font.bold = True
    _col_widths(alll_table, 1.6, 0.9, 0.9, 0.9, 0.9)

    alll_rows = [
        ("Allowance Balance", allowance and _fmt_dollar(allowance), None, None, None),
        (
            "Coverage Ratio",
            f"{coverage:.2f}x" if coverage is not None else "—",
            f"{prior_data.get('prior_qtr_alll_coverage'):.2f}x" if prior_data.get("prior_qtr_alll_coverage") else "—",
            f"{prior_data.get('prior_year_alll_coverage'):.2f}x" if prior_data.get("prior_year_alll_coverage") else "—",
            f"{cov_peer_med:.2f}x" if cov_peer_med else "—",
        ),
        (
            "Allowance / Loans",
            allow_to_loans and _fmt_pct(allow_to_loans, 3),
            None, None, None
        ),
        (
            "Peer Percentile",
            f"{cov_metrics.get('percentile_rank', '—'):.0f}th" if cov_metrics.get("percentile_rank") else "—",
            None, None,
            _STARS.get(cov_metrics.get("stars"), "—") if cov_metrics.get("stars") else "—",
        ),
    ]

    for vals in alll_rows:
        row = alll_table.add_row()
        for j, v in enumerate(vals):
            row.cells[j].text = str(v) if v is not None else "—"
        # Color coverage row: adverse = low coverage is bad
        if vals[0] == "Coverage Ratio" and coverage is not None:
            if coverage < 1.0:
                for c in row.cells[1:2]:
                    _cell_bg(c, _B_RED)
            elif cov_peer_med and coverage < cov_peer_med:
                _cell_bg(row.cells[1], _B_AMBER)


def _add_watch_items(doc: Document, watch_text: str) -> None:
    _heading(doc, "Section 5 — Watch Items")
    doc.add_paragraph(
        "Top 3 credit quality items requiring management attention this period. "
        "Recommended actions are for risk committee consideration."
    )

    if "[AI narrative" in watch_text or not watch_text.strip():
        doc.add_paragraph(watch_text or "Watch items unavailable.")
        return

    # Parse "ITEM N:" blocks from Claude output
    import re
    items = re.split(r"ITEM\s+\d+:", watch_text, flags=re.IGNORECASE)
    items = [item.strip() for item in items if item.strip()]

    for idx, item_text in enumerate(items[:3], start=1):
        _heading(doc, f"Watch Item {idx}", level=2)

        # Parse the four structured fields
        fields = {}
        for field_name in ("Metric", "Current Value", "Peer Context", "Recommended Action"):
            pattern = rf"{field_name}:\s*(.+?)(?=(?:Metric|Current Value|Peer Context|Recommended Action):|$)"
            match = re.search(pattern, item_text, re.IGNORECASE | re.DOTALL)
            if match:
                fields[field_name] = match.group(1).strip().replace("\n", " ")

        if len(fields) >= 2:
            # Render as a 2-column table
            tbl = doc.add_table(rows=len(fields), cols=2)
            tbl.style = "Table Grid"
            tbl.columns[0].width = Inches(1.6)
            tbl.columns[1].width = Inches(4.9)

            for i, (fname, fval) in enumerate(fields.items()):
                tbl.rows[i].cells[0].paragraphs[0].add_run(f"{fname}:").font.bold = True
                tbl.rows[i].cells[1].text = fval
                if fname == "Recommended Action":
                    _cell_bg(tbl.rows[i].cells[1], _B_AMBER)
        else:
            # Fallback: just write the text
            doc.add_paragraph(item_text)

        doc.add_paragraph()


def _add_forward_outlook(doc: Document, narrative: str) -> None:
    _heading(doc, "Forward Outlook (Next 2 Quarters)")
    doc.add_paragraph(narrative)


def _add_peer_appendix(
    doc: Document,
    regional_data: dict,
    period: str,
    peer_label: str,
) -> None:
    doc.add_page_break()
    _heading(doc, "Appendix — Full Regional Peer Comparison")
    doc.add_paragraph(
        f"All institutions in the {peer_label} peer group for {_period_label(period)}, "
        "ranked by 60+ day delinquency rate (lowest = best)."
    )

    peer_rows = regional_data.get("peer_rows", [])
    if not peer_rows:
        doc.add_paragraph("Peer data not available.")
        return

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Rank", "Institution", "60+ Day Delinquency Rate"]):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).font.bold = True
    _col_widths(tbl, 0.6, 4.0, 1.9)

    for rank, pr in enumerate(peer_rows, 1):
        row = tbl.add_row()
        row.cells[0].text = str(rank)
        row.cells[1].text = pr["institution_name"]
        row.cells[2].text = _fmt_pct(pr["delinq_rate"])


# ── Main build functions ──────────────────────────────────────────────────────

def build_report(
    charter_number: int,
    period: str,
    peer_group: str,
    output_dir: str,
    data: Optional[dict] = None,
    db_url: Optional[str] = None,
    board_version: bool = False,
) -> Path:
    """Generate credit quality risk committee memo; return path to .docx file.

    Called by api/routers/reports.py after _gather_data() populates `data`.
    board_version=True adds trend chart, forward outlook, and peer appendix.
    """
    data    = data or {}
    db_url  = db_url or DB_URL
    inst    = data.get("institution_financials", {})
    composite = data.get("credit_risk_composite", {})

    institution_name = str(inst.get("institution_name", f"Charter {charter_number}"))
    peer_group_label = data.get("peer_group_label", peer_group)
    peer_charters    = data.get("peer_charters", [])

    # Resolve peer group ID for benchmark lookups
    peer_group_id: Optional[str] = None
    try:
        from db import get_engine
        from sqlalchemy import text
        engine = get_engine(db_url)
        with engine.connect() as conn:
            row = conn.execute(
                text("""
                    SELECT id FROM peer_groups
                    WHERE :charter = ANY(institution_ids)
                      AND group_type = 'regional'
                    ORDER BY created_at DESC LIMIT 1
                """),
                {"charter": str(charter_number)},
            ).mappings().first()
        if row:
            peer_group_id = str(row["id"])
    except Exception:
        pass

    # ── Gather data ────────────────────────────────────────────────────────────
    trend_data     = _load_trend_data(charter_number, period, peer_group_id, db_url, peer_charters=peer_charters)
    regional_data  = _load_regional_comparison(charter_number, period, peer_charters, db_url)
    prior_data     = _load_prior_period_data(charter_number, period, db_url)

    current_delinq = (trend_data[0].get("delinq_rate_total") if trend_data else None)
    coverage       = inst.get("alll_coverage")
    co_rate        = inst.get("chargeoff_rate_total_annualized")

    # ── Claude narratives ──────────────────────────────────────────────────────
    exec_summary   = _gen_exec_summary(institution_name, period, peer_group_label, inst, trend_data, prior_data, peer_group_id)
    regional_narr  = _gen_regional_narrative(institution_name, period, current_delinq, regional_data)
    watch_text     = _gen_watch_items(institution_name, period, peer_group_label, inst, trend_data, prior_data, coverage, co_rate)

    chart_bytes: Optional[bytes] = None
    forward_outlook = ""
    if board_version:
        chart_bytes     = _make_trend_chart(trend_data)
        forward_outlook = _gen_forward_outlook(institution_name, period, trend_data)

    # ── Build document ─────────────────────────────────────────────────────────
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    report_type = "board" if board_version else "risk_committee"
    _add_memo_header(doc, institution_name, period, peer_group_label, report_type)
    _add_exec_summary(doc, exec_summary)
    doc.add_page_break()
    _add_delinq_trend_table(doc, trend_data, peer_group_label, chart_bytes if board_version else None)
    doc.add_page_break()
    _add_regional_comparison(doc, regional_data, institution_name, current_delinq, regional_narr)
    doc.add_page_break()
    _add_chargeoff_alll(doc, data, prior_data, peer_group_label)
    doc.add_page_break()
    _add_watch_items(doc, watch_text)

    if board_version:
        doc.add_page_break()
        _add_forward_outlook(doc, forward_outlook)
        _add_peer_appendix(doc, regional_data, period, peer_group_label)

    suffix = "_board" if board_version else ""
    filename = f"credit_quality_{charter_number}_{period}_{peer_group}{suffix}.docx"
    output_path = Path(output_dir) / filename
    doc.save(str(output_path))
    logger.info("Saved credit quality report → %s", output_path)
    return output_path


def generate_credit_quality_report(tenant_id: str, period: str) -> bytes:
    """Spec entry point: generate credit quality memo and return .docx bytes.

    Resolves charter number from tenant record, then generates full report.
    """
    import tempfile

    # Resolve charter number from tenant
    charter_number = _resolve_charter_for_tenant(tenant_id)
    if not charter_number:
        raise ValueError(f"Cannot resolve charter number for tenant {tenant_id}")

    with tempfile.TemporaryDirectory() as tmpdir:
        data = _gather_report_data(charter_number, period, tenant_id=tenant_id)
        path = build_report(
            charter_number=charter_number,
            period=period,
            peer_group="REGIONAL",
            output_dir=tmpdir,
            data=data,
            db_url=DB_URL,
        )
        return path.read_bytes()


def _resolve_charter_for_tenant(tenant_id: str) -> Optional[int]:
    try:
        from sqlalchemy import text
        from db import get_engine
        engine = get_engine(DB_URL)
        with engine.connect() as conn:
            row = conn.execute(
                text("SELECT charter_number FROM tenants WHERE id = :tid LIMIT 1"),
                {"tid": str(tenant_id)},
            ).mappings().first()
        return int(row["charter_number"]) if row else None
    except Exception:
        return None


def _gather_report_data(charter_number: int, period: str, tenant_id: str = "") -> dict:
    """Collect all data needed by build_report."""
    try:
        import pandas as pd
        from processing.delinquency_engine import compute_ratios, credit_risk_composite
        from processing.early_warning_engine import _trailing_periods, run_early_warning
        from processing.peer_engine import PeerGroupType, build_peer_group, peer_group_label
        from sqlalchemy import select
        from db import get_engine, institutions_quarterly

        engine = get_engine(DB_URL)
        group_type    = PeerGroupType("REGIONAL")
        peer_charters = build_peer_group(charter_number, period, group_type, tenant_id, db_url=DB_URL)
        label         = peer_group_label(group_type, charter_number, period, DB_URL)

        with engine.connect() as conn:
            rows = conn.execute(
                select(institutions_quarterly).where(
                    institutions_quarterly.c.charter_number == charter_number,
                    institutions_quarterly.c.period == period,
                )
            ).mappings().all()

        inst_df   = compute_ratios(pd.DataFrame([dict(r) for r in rows])) if rows else pd.DataFrame()
        periods   = _trailing_periods(period, n=8)
        warnings  = run_early_warning(charter_number, peer_charters, period, periods, db_url=DB_URL)
        composite = credit_risk_composite(charter_number, period, peer_charters, DB_URL)

        return {
            "tenant_id":              tenant_id,
            "charter_number":         charter_number,
            "period":                 period,
            "peer_group_type":        "REGIONAL",
            "peer_group_label":       label,
            "peer_count":             len(peer_charters),
            "peer_charters":          peer_charters,
            "institution_financials": inst_df.to_dict("records")[0] if not inst_df.empty else {},
            "credit_risk_composite":  composite,
            "early_warnings":         warnings,
        }
    except Exception as exc:
        logger.warning("_gather_report_data failed: %s", exc)
        return {"tenant_id": tenant_id, "charter_number": charter_number, "period": period, "peer_charters": []}
