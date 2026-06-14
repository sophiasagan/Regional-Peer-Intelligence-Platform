"""Quarterly competitive intelligence board report — python-docx + Claude AI.

Spec (reports/quarterly_template.py):
  Cover:     "[CU Name] Competitive Intelligence Report — Q[N] [YYYY]"
             "CONFIDENTIAL — For Internal Use Only"

  Section 1: Executive Summary         — Claude-generated, ~200 words
  Section 2: Market Share Dashboard    — one row per monitored geography
  Section 3: Competitive Movements     — Claude-generated per geography
  Section 4: Credit Quality Summary    — Callahan style, dollar balance + rate
  Section 5: Market Opportunities      — Claude-generated, top 3 geographies
  Section 6: Data Notes               — confidence tiers, sources, next releases

Returns: .docx bytes or Path depending on caller.
"""

from __future__ import annotations

import logging
import os
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
DB_URL            = os.environ.get("DATABASE_URL")

# ── Callahan / P76 color palette ──────────────────────────────────────────────
_C_GREEN  = RGBColor(0x2E, 0x7D, 0x32)   # text: top decile
_C_RED    = RGBColor(0xC6, 0x28, 0x28)   # text: bottom decile
_C_AMBER  = RGBColor(0xFF, 0x6F, 0x00)   # text: estimated / watch
_C_PURPLE = RGBColor(0x6A, 0x1B, 0x9A)   # text: regional peer label
_C_TEAL   = RGBColor(0x00, 0x69, 0x6E)   # text: measured confidence
_C_BLUE   = RGBColor(0x01, 0x57, 0x9B)   # text: modeled confidence

# Cell background shades (light tints)
_B_GREEN  = "E8F5E9"
_B_RED    = "FFEBEE"
_B_AMBER  = "FFF8E1"
_B_GRAY   = "F5F5F5"
_B_PURPLE = "F3E5F5"

_STARS = {1: "★☆☆☆☆", 2: "★★☆☆☆", 3: "★★★☆☆", 4: "★★★★☆", 5: "★★★★★"}


# ── python-docx helpers ───────────────────────────────────────────────────────

def _cell_bg(cell, hex6: str) -> None:
    """Set cell background color (hex without #)."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex6}" w:color="auto" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _bold_run(paragraph, text: str, color: Optional[RGBColor] = None) -> None:
    r = paragraph.add_run(text)
    r.font.bold = True
    if color:
        r.font.color.rgb = color


def _col_widths(table, *widths_in_inches) -> None:
    for i, w in enumerate(widths_in_inches):
        if i < len(table.columns):
            table.columns[i].width = Inches(w)


def _table_header(table, labels: list[str]) -> None:
    row = table.rows[0]
    for i, label in enumerate(labels):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        _bold_run(p, label)


def _fmt_pct(v: float, decimals: int = 2) -> str:
    return f"{v * 100:.{decimals}f}%"


def _fmt_dollar(v: float) -> str:
    if v >= 1e9:
        return f"${v / 1e9:.2f}B"
    if v >= 1e6:
        return f"${v / 1e6:.1f}M"
    if v >= 1e3:
        return f"${v / 1e3:.0f}K"
    return f"${v:,.0f}"


def _period_label(period: str) -> str:
    """'2026Q1' → 'Q1 2026'"""
    if "Q" in period:
        return f"Q{period[5]} {period[:4]}"
    return period


# ── Market share data helper ──────────────────────────────────────────────────

def _load_tenant_geographies(tenant_id: str, db_url: Optional[str]) -> list[dict]:
    """Return list of {geography_type, geography_id, group_name} for this tenant."""
    try:
        from sqlalchemy import text
        from db import get_engine
        engine = get_engine(db_url)
        with engine.connect() as conn:
            rows = conn.execute(
                text("""
                    SELECT group_name, geography_type, geography_ids
                    FROM peer_groups
                    WHERE tenant_id = :tid
                      AND geography_type IN ('county', 'msa', 'state')
                      AND geography_ids IS NOT NULL
                      AND array_length(geography_ids, 1) > 0
                    ORDER BY created_at DESC
                    LIMIT 10
                """),
                {"tid": str(tenant_id)},
            ).mappings().all()

        result = []
        for row in rows:
            ids = row["geography_ids"]
            if ids:
                result.append({
                    "geography_type": row["geography_type"],
                    "geography_id":   ids[0],    # primary ID for the group
                    "group_name":     row["group_name"] or ids[0],
                })
        return result
    except Exception:
        return []


def _fetch_market_share_row(
    charter_number: int,
    geo_type: str,
    geo_id: str,
    period: str,
    db_url: Optional[str],
) -> dict:
    """Return institution's market share data for one geography + period."""
    try:
        from processing.market_share_engine import calculate_market_share
        df = calculate_market_share(
            geography_type=geo_type,
            geography_id=geo_id,
            period=period,
            metric="deposits",
            institution_types=["bank", "cu"],
            db_url=db_url,
        )
        if df.empty:
            return {}

        inst_row = df[df["charter_or_cert"] == f"ncua:{charter_number}"]
        if inst_row.empty:
            return {}

        r = inst_row.iloc[0]
        total = df["metric_value"].sum()
        rank  = int((df["market_share"] > float(r["market_share"])).sum()) + 1

        # Top 3 competitors by share (excluding institution itself)
        competitors = (
            df[df["charter_or_cert"] != f"ncua:{charter_number}"]
            .nlargest(3, "market_share")[["institution_name", "market_share", "share_change_prior_period", "institution_type"]]
            .to_dict("records")
        )

        return {
            "your_share":                float(r["market_share"]),
            "share_change_prior_period": r.get("share_change_prior_period"),
            "share_change_yoy":          r.get("share_change_yoy"),
            "rank":                      rank,
            "n_institutions":            len(df),
            "total_market":              total,
            "confidence":                r.get("confidence", "modeled"),
            "data_period":               r.get("data_period", period),
            "competitors":               competitors,
        }
    except Exception:
        return {}


def _build_market_dashboard_data(
    charter_number: int,
    period: str,
    tenant_id: str,
    institution_state: str,
    db_url: Optional[str],
) -> list[dict]:
    """Assemble market share rows for the dashboard table."""
    geos = _load_tenant_geographies(tenant_id, db_url)

    # Fallback: institution's home state if no geographies configured
    if not geos:
        geos = [{"geography_type": "state", "geography_id": institution_state, "group_name": institution_state}]

    rows = []
    for geo in geos[:8]:   # cap at 8 geographies for readability
        share_data = _fetch_market_share_row(
            charter_number, geo["geography_type"], geo["geography_id"], period, db_url
        )
        if share_data:
            rows.append({**geo, **share_data})

    return rows


# ── Claude narrative generators ───────────────────────────────────────────────

def _claude(prompt: str, max_words: int = 250) -> str:
    """Call Claude and return the response text. Returns empty string on error."""
    if not ANTHROPIC_API_KEY:
        return "[AI narrative unavailable — ANTHROPIC_API_KEY not set]"
    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=max_words * 2,   # ~2 tokens/word
            system=(
                "You are a credit union board report writer. Write concise, professional "
                "narratives using Callahan Associates metric terminology. Be specific about "
                "numbers. Write in present tense for current conditions. No markdown."
            ),
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text.strip()
    except Exception as exc:
        logger.warning("Claude call failed: %s", exc)
        return "[AI narrative generation failed]"


def _generate_executive_summary(
    institution_name: str,
    period: str,
    peer_group_label: str,
    dashboard_rows: list[dict],
    inst_financials: dict,
    composite: dict,
) -> str:
    pctile = composite.get("composite_percentile")
    delinq = inst_financials.get("delinq_rate_total")
    co_rate = inst_financials.get("chargeoff_rate_total_annualized")
    assets = inst_financials.get("acct_010")
    members = inst_financials.get("acct_083")

    share_summary = ""
    if dashboard_rows:
        top_market = max(dashboard_rows, key=lambda r: r.get("your_share", 0))
        share_summary = (
            f"Highest market share: {top_market['group_name']} at "
            f"{top_market['your_share']*100:.1f}% (rank {top_market['rank']} of "
            f"{top_market['n_institutions']}). "
        )
        gainers = [r for r in dashboard_rows if (r.get("share_change_prior_period") or 0) > 0]
        if gainers:
            share_summary += f"Share increased QoQ in {len(gainers)} monitored market(s). "

    prompt = (
        f"Write a 200-word executive summary for {institution_name}'s board competitive "
        f"intelligence report for {_period_label(period)}. "
        f"Peer group: {peer_group_label}. "
        f"{'Total assets: ' + _fmt_dollar(assets) + '. ' if assets else ''}"
        f"{'Members: ' + f'{members:,.0f}. ' if members else ''}"
        f"{'Credit quality composite: ' + f'{pctile:.0f}th percentile vs peers. ' if pctile else ''}"
        f"{'Total delinquency ratio: ' + _fmt_pct(delinq) + '. ' if delinq else ''}"
        f"{'Net charge-off ratio (annualized): ' + _fmt_pct(co_rate) + '. ' if co_rate else ''}"
        f"{share_summary}"
        "Summarize market position, key movements this quarter, and the single most "
        "important competitive opportunity. Be specific and board-appropriate."
    )
    return _claude(prompt, max_words=220)


def _generate_competitive_movements(
    institution_name: str,
    period: str,
    geo: dict,
) -> str:
    competitors = geo.get("competitors", [])
    if not competitors:
        return "Competitor data not available for this geography."

    comp_text = "; ".join(
        f"{c['institution_name']} ({c['institution_type'].upper()}) at "
        f"{c['market_share']*100:.1f}% share"
        + (f", {'+' if (c.get('share_change_prior_period') or 0) >= 0 else ''}"
           f"{(c.get('share_change_prior_period') or 0)*100:.2f} pp QoQ" if c.get("share_change_prior_period") else "")
        for c in competitors
    )

    your_share = geo.get("your_share", 0)
    rank = geo.get("rank", "?")
    n    = geo.get("n_institutions", "?")

    prompt = (
        f"Write a 120-word competitive movements narrative for {institution_name} in "
        f"{geo.get('group_name', 'the monitored market')} for {_period_label(period)}. "
        f"Our institution holds {your_share*100:.1f}% deposit market share (rank {rank} of {n}). "
        f"Top competitors: {comp_text}. "
        "Describe who gained or lost share, any notable movements, and what this means "
        "for competitive positioning. Be specific and actionable."
    )
    return _claude(prompt, max_words=140)


def _generate_market_opportunities(
    institution_name: str,
    period: str,
    institution_state: str,
    dashboard_rows: list[dict],
) -> str:
    low_penetration = sorted(dashboard_rows, key=lambda r: r.get("your_share", 1))[:3]
    summary = "; ".join(
        f"{r.get('group_name')} (our share: {r.get('your_share', 0)*100:.1f}%, "
        f"rank {r.get('rank', '?')} of {r.get('n_institutions', '?')})"
        for r in low_penetration
    )

    prompt = (
        f"Write a 180-word market opportunities section for {institution_name}'s board report, "
        f"{_period_label(period)}, based in {institution_state}. "
        f"Geographies with growth opportunity (lowest current share): {summary}. "
        "Identify the top 3 geographic expansion opportunities based on: (1) low credit union "
        "penetration relative to total market, (2) competitive weakness or market share gaps, "
        "(3) demographic alignment with credit union membership growth. "
        "Format as three numbered opportunity statements. Be specific."
    )
    return _claude(prompt, max_words=200)


# ── Section builders ──────────────────────────────────────────────────────────

def _add_cover_page(doc: Document, institution_name: str, period: str, peer_group_label: str) -> None:
    # Vertical padding
    for _ in range(4):
        doc.add_paragraph()

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{institution_name}\nCompetitive Intelligence Report — {_period_label(period)}")
    run.font.size = Pt(22)
    run.font.bold = True

    doc.add_paragraph()

    # Confidential line
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = conf.add_run("CONFIDENTIAL — For Internal Use Only")
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = _C_RED

    doc.add_paragraph()

    # Metadata block
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Period: {_period_label(period)}  ·  "
        f"Peer Group: {peer_group_label}  ·  "
        f"Generated: {date.today().strftime('%B %d, %Y')}"
    ).font.size = Pt(10)

    doc.add_page_break()


def _add_executive_summary(doc: Document, narrative: str) -> None:
    _heading(doc, "Section 1 — Executive Summary")
    doc.add_paragraph(narrative)


def _add_market_share_dashboard(doc: Document, dashboard_rows: list[dict], period: str) -> None:
    _heading(doc, "Section 2 — Market Share Dashboard")
    doc.add_paragraph(
        f"Deposit market share by monitored geography — {_period_label(period)}. "
        "Green: share improved vs prior quarter. Amber: share declined. "
        "All figures include both credit unions and banks."
    )

    if not dashboard_rows:
        doc.add_paragraph("No market share data available. Ensure cu_deposit_allocations and fdic_deposits tables are populated.")
        return

    cols = ["Geography", "Type", "Your Share", "Prior Qtr", "Δ QoQ", "Δ YoY", "Rank", "Market Total", "Confidence"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    _table_header(table, cols)
    _col_widths(table, 1.3, 0.5, 0.8, 0.8, 0.7, 0.7, 0.5, 1.0, 0.9)

    for r in dashboard_rows:
        row = table.add_row()
        share  = r.get("your_share", 0)
        qoq    = r.get("share_change_prior_period")
        yoy    = r.get("share_change_yoy")

        row.cells[0].text = r.get("group_name", r.get("geography_id", "?"))
        row.cells[1].text = r.get("geography_type", "").upper()
        row.cells[2].text = _fmt_pct(share, 2) if share else "N/A"
        prior = share - qoq if (qoq is not None and share) else None
        row.cells[3].text = _fmt_pct(prior, 2) if prior else "N/A"
        row.cells[4].text = (f"{'+' if (qoq or 0) >= 0 else ''}{qoq*100:.2f} pp") if qoq is not None else "N/A"
        row.cells[5].text = (f"{'+' if (yoy or 0) >= 0 else ''}{yoy*100:.2f} pp") if yoy is not None else "N/A"
        row.cells[6].text = f"{r.get('rank', '?')} / {r.get('n_institutions', '?')}"
        tot = r.get("total_market")
        row.cells[7].text = _fmt_dollar(tot) if tot else "N/A"
        row.cells[8].text = r.get("confidence", "modeled").title()

        # Green/amber row shading based on QoQ direction
        bg = _B_GREEN if (qoq is not None and qoq > 0) else _B_AMBER if (qoq is not None and qoq < 0) else _B_GRAY
        for cell in row.cells:
            _cell_bg(cell, bg)

    # Confidence legend
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Confidence tiers: ")
    r = p.add_run("Measured "); r.font.color.rgb = _C_TEAL
    p.add_run("(FDIC branch data)  ")
    r = p.add_run("Modeled "); r.font.color.rgb = _C_BLUE
    p.add_run("(CU allocation model, ±8% validated)  ")
    r = p.add_run("Estimated "); r.font.color.rgb = _C_AMBER
    p.add_run("(proxy-based)")


def _add_competitive_movements(doc: Document, dashboard_rows: list[dict], narratives: dict[str, str]) -> None:
    _heading(doc, "Section 3 — Competitive Movements")
    doc.add_paragraph(
        "Identifies which competitors gained or lost market share in each monitored geography "
        "and interprets what this means for competitive positioning."
    )

    if not dashboard_rows:
        doc.add_paragraph("No competitive movement data available.")
        return

    for geo in dashboard_rows:
        group_name = geo.get("group_name", geo.get("geography_id", "?"))
        _heading(doc, group_name, level=2)

        # Competitor table
        competitors = geo.get("competitors", [])
        if competitors:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            _table_header(table, ["Institution", "Type", "Share %", "QoQ Change"])
            _col_widths(table, 2.5, 0.6, 0.8, 1.0)

            for c in competitors:
                row = table.add_row()
                row.cells[0].text = c.get("institution_name", "?")
                row.cells[1].text = (c.get("institution_type") or "?").upper()
                row.cells[2].text = _fmt_pct(c.get("market_share", 0))
                qoq = c.get("share_change_prior_period")
                row.cells[3].text = (f"{'+' if (qoq or 0) >= 0 else ''}{qoq*100:.2f} pp") if qoq is not None else "N/A"

                bg = _B_RED if (qoq and qoq < 0) else _B_GREEN if (qoq and qoq > 0) else _B_GRAY
                for cell in row.cells:
                    _cell_bg(cell, bg)

            doc.add_paragraph()

        # AI narrative
        narrative = narratives.get(group_name, "")
        if narrative:
            doc.add_paragraph(narrative)


def _add_credit_quality_section(doc: Document, data: dict, peer_group_label: str) -> None:
    _heading(doc, "Section 4 — Credit Quality Summary")
    doc.add_paragraph(
        f"Key credit quality metrics vs {peer_group_label}. "
        "Callahan star rating: ★★★★★ = top 10% (best), ★☆☆☆☆ = bottom 10% (worst). "
        "Delinquency is an ADVERSE metric — lower rate = better = more stars. "
        "All delinquency figures are institution-level (not branch-level)."
    )

    inst = data.get("institution_financials", {})
    composite = data.get("credit_risk_composite", {})
    comp_metrics = composite.get("metrics", {})

    # ── Delinquency and charge-off metrics with both dollar balance AND rate ──
    # CLAUDE.md rule: "Always display BOTH dollar balance AND computed rate"
    metrics = [
        # (label, balance_key, rate_key, is_adverse)
        ("Total 60+ Day Delinquency",    "acct_041B",   "delinq_rate_total",               True),
        ("90+ Day Delinquency",          None,          "delinq_rate_90plus",               True),
        ("Gross Charge-Offs YTD",        "acct_550",    "chargeoff_rate_total_annualized",  True),
        ("ACL / ALLL on Loans",          "acct_AS0048", "alll_to_loans",                    False),
        ("Allowance Coverage Ratio",     None,          "alll_coverage",                    False),
        ("Net Worth Ratio",              "acct_797",    "net_worth_ratio",                  False),
        ("Return on Assets (Ann.)",      None,          "roa_annualized",                   False),
        ("Efficiency Ratio",             None,          "efficiency_ratio",                 True),
    ]

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    _table_header(table, ["Metric", "Balance ($)", "Your Rate", "Peer Median", "Δ vs Peer", "Stars", "Pctile"])
    _col_widths(table, 2.0, 1.0, 0.9, 1.0, 0.8, 0.8, 0.6)

    for label, bal_key, rate_key, is_adverse in metrics:
        row = table.add_row()
        row.cells[0].text = label

        # Balance
        bal = inst.get(bal_key) if bal_key else None
        if bal_key == "acct_AS0048":
            bal = bal or inst.get("acct_719")   # CECL → pre-CECL fallback
        row.cells[1].text = _fmt_dollar(bal) if bal else "—"

        # Rate
        rate = inst.get(rate_key)
        row.cells[2].text = (
            f"{rate:.2f}x" if rate_key == "alll_coverage" and rate is not None
            else _fmt_pct(rate, 3) if rate is not None
            else "—"
        )

        # Peer stats from composite
        m = comp_metrics.get(rate_key, {})
        peer_dist = m.get("peer_distribution", {})
        peer_med = peer_dist.get("p50")
        pct_rank = m.get("percentile_rank")
        stars    = m.get("stars")

        row.cells[3].text = (
            f"{peer_med:.2f}x" if rate_key == "alll_coverage" and peer_med is not None
            else _fmt_pct(peer_med, 3) if peer_med is not None
            else "—"
        )

        # Delta vs peer median (how many ppt above/below)
        if rate is not None and peer_med is not None:
            delta = rate - peer_med
            sign  = "+" if delta >= 0 else ""
            row.cells[4].text = (
                f"{sign}{delta:.2f}x" if rate_key == "alll_coverage"
                else f"{sign}{delta*100:.2f} pp"
            )
            # Adverse: positive delta is bad (red); negative is good (green)
            # Positive: positive delta is good (green); negative is bad (red)
            good = (delta < 0) if is_adverse else (delta > 0)
            _cell_bg(row.cells[4], _B_GREEN if good else _B_RED)
        else:
            row.cells[4].text = "—"

        row.cells[5].text = _STARS.get(stars, "—") if stars else "—"
        row.cells[6].text = f"{pct_rank:.0f}th" if pct_rank is not None else "—"

        # Color star column by decile (Callahan convention)
        if pct_rank is not None:
            if pct_rank >= 90:
                _cell_bg(row.cells[5], _B_GREEN)
                row.cells[5].paragraphs[0].runs[0].font.color.rgb = _C_GREEN
            elif pct_rank < 10:
                _cell_bg(row.cells[5], _B_RED)
                row.cells[5].paragraphs[0].runs[0].font.color.rgb = _C_RED

    # Early warning summary appended below the table
    warnings = data.get("early_warnings", [])
    active_warnings = [w for w in warnings if getattr(getattr(w, "alert_level", None), "value", "none") not in ("none", "green")]
    if active_warnings:
        doc.add_paragraph()
        _heading(doc, "Know Before Your Examiner Does", level=2)
        for w in active_warnings:
            level_val = getattr(getattr(w, "alert_level", None), "value", "watch")
            p    = doc.add_paragraph()
            badge = p.add_run(f"[{level_val.upper()}] ")
            badge.font.bold = True
            badge.font.color.rgb = _C_RED if level_val in ("red", "urgent") else _C_AMBER
            p.add_run(w.message)


def _add_market_opportunities(doc: Document, narrative: str) -> None:
    _heading(doc, "Section 5 — Market Opportunities")
    doc.add_paragraph(
        "Geographies where market conditions favor growth: low credit union penetration, "
        "demographic alignment, or competitor weakness. Ranked by estimated opportunity."
    )
    if narrative:
        doc.add_paragraph(narrative)
    else:
        doc.add_paragraph("Insufficient market share data to generate opportunity analysis. "
                          "Ensure cu_deposit_allocations and fdic_deposits tables are populated.")


def _add_data_notes(
    doc: Document,
    dashboard_rows: list[dict],
    period: str,
    peer_group_label: str,
) -> None:
    _heading(doc, "Section 6 — Data Notes")

    # Confidence summary
    conf_counts: dict[str, int] = {"measured": 0, "modeled": 0, "estimated": 0}
    for row in dashboard_rows:
        c = row.get("confidence", "estimated")
        conf_counts[c] = conf_counts.get(c, 0) + 1

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _table_header(table, ["Confidence Level", "Geographies", "Description"])
    _col_widths(table, 1.3, 1.0, 4.0)

    conf_info = [
        ("Measured",  "measured",  _C_TEAL,   "FDIC Summary of Deposits (branch-level) — highest confidence"),
        ("Modeled",   "modeled",   _C_BLUE,   "CU deposit allocation model, ±8% cross-validated vs FDIC totals"),
        ("Estimated", "estimated", _C_AMBER,  "Proxy-based geographic allocation — flag for management attention"),
    ]
    for display, key, color, desc in conf_info:
        row = table.add_row()
        r = row.cells[0].paragraphs[0].add_run(display)
        r.font.bold  = True
        r.font.color.rgb = color
        row.cells[1].text = str(conf_counts.get(key, 0))
        row.cells[2].text = desc

    # Data sources table
    doc.add_paragraph()
    _heading(doc, "Data Sources", level=2)

    sources_table = doc.add_table(rows=1, cols=4)
    sources_table.style = "Table Grid"
    _table_header(sources_table, ["Source", "Data Type", "Period Used", "Next Release"])
    _col_widths(sources_table, 1.8, 1.8, 1.2, 1.8)

    year = int(period[:4])
    quarter = int(period[5]) if "Q" in period else 4

    # Static release schedule (approximate)
    ncua_next = f"~6 weeks after Q{(quarter % 4) + 1} {year if quarter < 4 else year + 1} end"
    fdic_next  = f"~August {year + 1} (for {year} Summary of Deposits)"
    hmda_next  = f"~March {year + 1} (for {year} HMDA LAR)"

    sources = [
        ("NCUA 5300 Call Report", "Credit quality, loans, shares, members", _period_label(period), ncua_next),
        ("FDIC Summary of Deposits", "Bank deposit market share (branch-level)", str(year), fdic_next),
        ("HMDA Loan Application Register", "Mortgage origination market share", str(year - 1), hmda_next),
        ("Census ACS 5-Year", "Demographics, market sizing", f"{year - 2}–{year}", "Annual (December)"),
    ]

    for name, data_type, period_used, next_rel in sources:
        row = sources_table.add_row()
        row.cells[0].text = name
        row.cells[1].text = data_type
        row.cells[2].text = period_used
        row.cells[3].text = next_rel

    # Disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Note: ")
    r.font.bold = True
    p.add_run(
        "Market share figures for credit union geographies are derived from an allocation model "
        "applied to NCUA institution-level data. Modeled figures are validated at ±8% against "
        "available FDIC branch data. Estimated figures are based on proxy methods and should be "
        "treated as indicative only. NCUA delinquency figures are institution-level (not "
        "branch-level) — confidence is always 'Measured' for credit quality metrics."
    )
    p.runs[-1].font.size = Pt(9)


# ── Main build function ───────────────────────────────────────────────────────

def build_report(
    charter_number: int,
    period: str,
    peer_group: str,
    output_dir: str,
    data: Optional[dict] = None,
    db_url: Optional[str] = None,
) -> Path:
    """Generate quarterly board report; return path to written .docx file.

    Called by api/routers/reports.py after _gather_data() populates `data`.
    """
    data    = data or {}
    db_url  = db_url or DB_URL
    inst    = data.get("institution_financials", {})
    composite = data.get("credit_risk_composite", {})

    institution_name = (
        str(inst.get("institution_name", ""))
        or f"Charter {charter_number}"
    )
    institution_state  = str(inst.get("state_code", ""))
    peer_group_label   = data.get("peer_group_label", peer_group)
    tenant_id          = data.get("tenant_id", "")

    # ── Gather market share dashboard data ─────────────────────────────────────
    dashboard_rows = _build_market_dashboard_data(
        charter_number, period, tenant_id, institution_state, db_url
    )

    # ── Claude narrative generation (all at once to parallelize I/O if needed) ─
    exec_summary = _generate_executive_summary(
        institution_name, period, peer_group_label, dashboard_rows, inst, composite
    )

    comp_narratives: dict[str, str] = {}
    for geo in dashboard_rows:
        name = geo.get("group_name", geo.get("geography_id", "?"))
        comp_narratives[name] = _generate_competitive_movements(institution_name, period, geo)

    opportunities = _generate_market_opportunities(
        institution_name, period, institution_state, dashboard_rows
    )

    # ── Build document ─────────────────────────────────────────────────────────
    doc = Document()

    # Set default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    _add_cover_page(doc, institution_name, period, peer_group_label)
    _add_executive_summary(doc, exec_summary)
    doc.add_page_break()
    _add_market_share_dashboard(doc, dashboard_rows, period)
    doc.add_page_break()
    _add_competitive_movements(doc, dashboard_rows, comp_narratives)
    doc.add_page_break()
    _add_credit_quality_section(doc, data, peer_group_label)
    doc.add_page_break()
    _add_market_opportunities(doc, opportunities)
    doc.add_page_break()
    _add_data_notes(doc, dashboard_rows, period, peer_group_label)

    filename = f"quarterly_{charter_number}_{period}_{peer_group}.docx"
    output_path = Path(output_dir) / filename
    doc.save(str(output_path))
    logger.info("Saved quarterly report → %s", output_path)
    return output_path


def generate_quarterly_report(tenant_id: str, period: str) -> bytes:
    """Spec entry point: generate report and return .docx bytes.

    Looks up the institution charter number from the tenant record.
    Falls back to temporary directory for file I/O.
    """
    import tempfile

    # Resolve charter_number from tenant_id
    charter_number = _resolve_charter_for_tenant(tenant_id)
    if not charter_number:
        raise ValueError(f"Cannot resolve charter number for tenant {tenant_id}")

    with tempfile.TemporaryDirectory() as tmpdir:
        # Build data payload (mirrors what reports.py _gather_data does)
        data = _gather_report_data(charter_number, period, tenant_id=tenant_id)
        path = build_report(
            charter_number=charter_number,
            period=period,
            peer_group="REGIONAL",
            output_dir=tmpdir,
            data=data,
            db_url=DB_URL,
        )
        return path.read_bytes()


def _resolve_charter_for_tenant(tenant_id: str) -> Optional[int]:
    """Look up primary charter number for a tenant_id from the tenants table."""
    try:
        from sqlalchemy import text
        from db import get_engine
        engine = get_engine(DB_URL)
        with engine.connect() as conn:
            row = conn.execute(
                text("SELECT charter_number FROM tenants WHERE id = :tid LIMIT 1"),
                {"tid": str(tenant_id)},
            ).mappings().first()
        return int(row["charter_number"]) if row else None
    except Exception:
        return None


def _gather_report_data(charter_number: int, period: str, tenant_id: str = "") -> dict:
    """Collect all data needed by build_report — mirrors reports.py._gather_data."""
    try:
        import pandas as pd
        from processing.delinquency_engine import compute_ratios, credit_risk_composite
        from processing.early_warning_engine import _trailing_periods, run_early_warning
        from processing.peer_engine import PeerGroupType, build_peer_group, peer_group_label
        from sqlalchemy import select
        from db import get_engine, institutions_quarterly

        engine = get_engine(DB_URL)
        group_type    = PeerGroupType("REGIONAL")
        peer_charters = build_peer_group(charter_number, period, group_type, tenant_id, db_url=DB_URL)
        label         = peer_group_label(group_type, charter_number, period, DB_URL)

        with engine.connect() as conn:
            result = conn.execute(
                select(institutions_quarterly).where(
                    institutions_quarterly.c.charter_number == charter_number,
                    institutions_quarterly.c.period == period,
                )
            )
            rows = result.mappings().all()

        inst_df   = compute_ratios(pd.DataFrame([dict(r) for r in rows])) if rows else pd.DataFrame()
        periods   = _trailing_periods(period, n=8)
        warnings  = run_early_warning(charter_number, peer_charters, period, periods, db_url=DB_URL)
        composite = credit_risk_composite(charter_number, period, peer_charters, DB_URL)

        return {
            "tenant_id":              tenant_id,
            "charter_number":         charter_number,
            "period":                 period,
            "peer_group_type":        "REGIONAL",
            "peer_group_label":       label,
            "peer_count":             len(peer_charters),
            "institution_financials": inst_df.to_dict("records")[0] if not inst_df.empty else {},
            "credit_risk_composite":  composite,
            "early_warnings":         warnings,
        }
    except Exception as exc:
        logger.warning("_gather_report_data failed: %s", exc)
        return {"tenant_id": tenant_id, "charter_number": charter_number, "period": period}
